"""
PDF text extraction tests.
"""

from __future__ import annotations

from io import BytesIO

import pandas as pd
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

from services.document_processor import DocumentProcessor
from services.report_markdown_renderer import parse_markdown_blocks


def _csv_upload(df: pd.DataFrame, name: str = "data.csv") -> BytesIO:
    buffer = BytesIO()
    df.to_csv(buffer, index=False)
    buffer.seek(0)
    buffer.name = name
    return buffer


def _xlsx_upload(sheets: dict[str, pd.DataFrame], name: str = "data.xlsx") -> BytesIO:
    buffer = BytesIO()
    with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
        for sheet_name, df in sheets.items():
            df.to_excel(writer, sheet_name=sheet_name, index=False)
    buffer.seek(0)
    buffer.name = name
    return buffer


def _sample_pdf_bytes(text: str = "Annual market revenue increased 12 percent.") -> bytes:
    buffer = BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=letter)
    pdf.drawString(72, 720, text)
    pdf.save()
    return buffer.getvalue()


def test_extract_pdf_text_from_bytes():
    data = _sample_pdf_bytes()
    text = DocumentProcessor._extract_pdf_text(data)

    assert "revenue" in text.lower()


def test_extract_pdf_from_upload_object():
    buffer = BytesIO(_sample_pdf_bytes("Board report summary for Q4."))
    buffer.name = "report.pdf"

    text = DocumentProcessor.extract_text(buffer)

    assert "board report" in text.lower()


def test_pdf_ocr_fallback_used_when_text_layer_empty(monkeypatch):
    data = _sample_pdf_bytes("ignored")

    monkeypatch.setattr(
        DocumentProcessor,
        "_extract_pdf_with_pypdf2",
        staticmethod(lambda *_args, **_kwargs: ""),
    )
    monkeypatch.setattr(
        DocumentProcessor,
        "_extract_pdf_with_pymupdf",
        staticmethod(lambda *_args, **_kwargs: ""),
    )
    monkeypatch.setattr(
        DocumentProcessor,
        "_extract_pdf_with_pdfplumber",
        staticmethod(lambda *_args, **_kwargs: ""),
    )
    monkeypatch.setattr(
        DocumentProcessor,
        "_extract_pdf_with_ocr",
        staticmethod(lambda *_args, **_kwargs: "OCR market revenue increased 12 percent."),
    )

    text = DocumentProcessor._extract_pdf_text(data)

    assert "OCR market revenue" in text


def test_pdf_table_extraction_runs_even_on_pages_with_a_text_layer(monkeypatch):
    """Report Output Quality Upgrade Step D: real PDFs (e.g. a NAICOM-style
    financial report) commonly have both a text layer AND genuine embedded
    tables on the same page. Before this fix, extract_tables() only ran on
    pages with NO text layer at all, so real tabular data in text-layered
    PDFs never reached quantitative_analysis_service.extract_metric_tables()
    and the dashboard fell back to generic, unlabeled charts."""

    import pdfplumber

    class _FakePage:
        def extract_text(self):
            return "Gross premium grew steadily across the reporting period."

        def extract_tables(self):
            return [
                [
                    ["Year", "Gross Premium"],
                    ["2022", "789.6"],
                    ["2023", "1043.1"],
                    ["2024", "1558.7"],
                ]
            ]

    class _FakeDocument:
        def __enter__(self):
            return self

        def __exit__(self, *_args):
            return False

        @property
        def pages(self):
            return [_FakePage()]

    monkeypatch.setattr(pdfplumber, "open", lambda *_args, **_kwargs: _FakeDocument())

    text = DocumentProcessor._extract_pdf_with_pdfplumber(b"fake-pdf-bytes")

    assert "Gross premium grew steadily" in text
    assert "| Year | Gross Premium |" in text
    assert "| 2023 | 1043.1 |" in text

    blocks = parse_markdown_blocks(text)
    tables = [block for block in blocks if block.block_type == "table"]
    assert len(tables) == 1
    assert tables[0].rows[0] == ["Year", "Gross Premium"]

    from services.quantitative_analysis_service import extract_metric_tables

    metric_tables = extract_metric_tables([{"filename": "report.pdf", "excerpt": text}])
    assert len(metric_tables) == 1
    assert metric_tables[0]["title"] == "Gross Premium"


def test_extract_pdf_text_includes_pdfplumber_tables_even_when_pypdf2_wins(monkeypatch):
    """The critical integration point: _extract_pdf_text() is what
    extract_text()/upload processing actually calls. Before this fix,
    pdfplumber's table-aware extractor was only ever reached as a
    plain-text fallback candidate — pypdf2 almost always wins the
    "first extractor to produce >=200 chars" race for a real PDF with a
    text layer, so pdfplumber (and its table detection) never ran at
    all. The table supplement must run regardless of which extractor
    won the narrative-text race."""

    import pdfplumber

    monkeypatch.setattr(
        DocumentProcessor,
        "_extract_pdf_with_pypdf2",
        staticmethod(
            lambda *_a, **_kw: "Gross premium grew steadily across the reporting period. " * 5
        ),
    )
    monkeypatch.setattr(
        DocumentProcessor, "_extract_pdf_with_pymupdf", staticmethod(lambda *_a, **_kw: "")
    )

    class _FakePage:
        def extract_text(self):
            return "Gross premium grew steadily across the reporting period."

        def extract_tables(self):
            return [[["Year", "Gross Premium"], ["2022", "789.6"], ["2023", "1043.1"]]]

    class _FakeDocument:
        def __enter__(self):
            return self

        def __exit__(self, *_args):
            return False

        @property
        def pages(self):
            return [_FakePage()]

    monkeypatch.setattr(pdfplumber, "open", lambda *_args, **_kwargs: _FakeDocument())

    text = DocumentProcessor._extract_pdf_text(b"fake-pdf-bytes")

    assert "Gross premium grew steadily" in text
    assert "| Year | Gross Premium |" in text
    assert "| 2023 | 1043.1 |" in text


def test_pdf_table_extraction_skips_a_single_row_table_fragment(monkeypatch):
    """A one-row detection (no real data rows) shouldn't be emitted as a
    table block — it can't support any calculation and would just be
    prose-like noise in the retrieved text."""

    import pdfplumber

    class _FakePage:
        def extract_text(self):
            return ""

        def extract_tables(self):
            return [[["Just one row"]]]

    class _FakeDocument:
        def __enter__(self):
            return self

        def __exit__(self, *_args):
            return False

        @property
        def pages(self):
            return [_FakePage()]

    monkeypatch.setattr(pdfplumber, "open", lambda *_args, **_kwargs: _FakeDocument())

    text = DocumentProcessor._extract_pdf_with_pdfplumber(b"fake-pdf-bytes")

    assert "Just one row" not in text


def test_extract_csv_includes_verified_summary_statistics():
    df = pd.DataFrame({"Region": ["West", "East", "North"], "Sales": [1100000, 900000, 500000]})
    text = DocumentProcessor.extract_text(_csv_upload(df))

    assert "count=3" in text
    assert f"sum={2500000:,}" in text
    assert f"mean={int(2500000 / 3):,}" in text
    assert f"min={500000:,}" in text
    assert f"max={1100000:,}" in text


def test_extract_csv_markdown_table_round_trips_through_report_renderer():
    df = pd.DataFrame({"Region": ["West", "East"], "Sales": [1100000, 900000]})
    text = DocumentProcessor.extract_text(_csv_upload(df))

    blocks = parse_markdown_blocks(text)
    tables = [block for block in blocks if block.block_type == "table"]

    assert len(tables) == 1
    assert tables[0].rows[0] == ["Region", "Sales"]
    assert tables[0].rows[1] == ["West", "1,100,000"]
    assert tables[0].rows[2] == ["East", "900,000"]


def test_extract_csv_pipe_character_in_cell_does_not_corrupt_columns():
    df = pd.DataFrame({"Region": ["West|East combined", "North"], "Sales": [100, 200]})
    text = DocumentProcessor.extract_text(_csv_upload(df))

    blocks = parse_markdown_blocks(text)
    tables = [block for block in blocks if block.block_type == "table"]

    assert len(tables) == 1
    assert all(len(row) == 2 for row in tables[0].rows)
    assert "|" not in tables[0].rows[1][0]


def test_extract_csv_truncation_preserves_full_dataset_stats():
    df = pd.DataFrame({"Amount": list(range(1, 21))})  # 20 rows, sum=210
    text = DocumentProcessor.extract_text(_csv_upload(df), max_tabular_rows=5)

    assert f"sum={210:,}" in text
    assert "count=20" in text

    blocks = parse_markdown_blocks(text)
    tables = [block for block in blocks if block.block_type == "table"]
    assert len(tables) == 1
    assert len(tables[0].rows) - 1 == 5  # header + 5 data rows


def test_extract_csv_no_row_cap_by_default():
    df = pd.DataFrame({"Amount": list(range(1, 21))})
    text = DocumentProcessor.extract_text(_csv_upload(df))

    blocks = parse_markdown_blocks(text)
    tables = [block for block in blocks if block.block_type == "table"]
    assert len(tables[0].rows) - 1 == 20


def test_extract_xlsx_all_sheets_present():
    sheets = {
        "Summary": pd.DataFrame({"Total": [4200000]}),
        "Regional": pd.DataFrame({"Region": ["West", "East"], "Amount": [1100000, 900000]}),
    }
    text = DocumentProcessor.extract_text(_xlsx_upload(sheets))

    assert "Sheet: Summary" in text
    assert "Sheet: Regional" in text
    assert "4,200,000" in text
    assert "West" in text and "East" in text


def test_extract_xlsx_empty_sheet_skipped_without_crashing():
    sheets = {
        "Data": pd.DataFrame({"Amount": [100]}),
        "Blank": pd.DataFrame(),
    }
    text = DocumentProcessor.extract_text(_xlsx_upload(sheets))

    # Only one sheet survives the empty-sheet filter, so it's treated like
    # a single-sheet workbook (no "Sheet: Data" heading, matching how a
    # plain CSV has no sheet heading either) — assert on its content.
    assert "sum=100" in text
    assert "Blank" not in text


def test_extract_csv_all_text_columns_states_no_numeric_data():
    df = pd.DataFrame({"Name": ["Alice", "Bob"], "Role": ["Engineer", "Manager"]})
    text = DocumentProcessor.extract_text(_csv_upload(df))

    assert "No numeric columns detected" in text


def test_extract_csv_empty_dataframe_headers_only():
    df = pd.DataFrame({"Region": pd.Series(dtype="object"), "Amount": pd.Series(dtype="float64")})
    text = DocumentProcessor.extract_text(_csv_upload(df))

    # An empty CSV gives pandas nothing to infer numeric dtypes from on
    # read-back, so "no numeric columns" is the correct, graceful outcome
    # here — not a crash, and not a fabricated stats block.
    assert "No numeric columns detected" in text
    assert "(no data rows)" in text


def test_extract_csv_all_nan_numeric_column_states_no_non_null_values():
    df = pd.DataFrame({"Region": ["West", "East"], "Amount": [None, None]})
    text = DocumentProcessor.extract_text(_csv_upload(df))

    assert "no non-null values" in text
    assert "nan" not in text.lower()


def test_extract_csv_nan_values_do_not_leak_literal_nan_text():
    df = pd.DataFrame({"Region": ["West", None], "Amount": [100, None]})
    text = DocumentProcessor.extract_text(_csv_upload(df))

    assert "nan" not in text.lower()


def test_extract_text_from_path_xlsx_multi_sheet_delegates_correctly(tmp_path):
    sheets = {
        "Summary": pd.DataFrame({"Total": [4200000]}),
        "Regional": pd.DataFrame({"Region": ["West"], "Amount": [1100000]}),
    }
    path = tmp_path / "workbook.xlsx"
    with pd.ExcelWriter(path, engine="openpyxl") as writer:
        for sheet_name, df in sheets.items():
            df.to_excel(writer, sheet_name=sheet_name, index=False)

    text = DocumentProcessor.extract_text_from_path(path)

    assert "Sheet: Summary" in text
    assert "Sheet: Regional" in text


def test_extract_txt_unaffected():
    buffer = BytesIO("Board meeting minutes for Q3.".encode("utf-8"))
    buffer.name = "notes.txt"

    text = DocumentProcessor.extract_text(buffer)

    assert text == "Board meeting minutes for Q3."


def test_extract_docx_unaffected():
    from docx import Document as DocxDocument

    document = DocxDocument()
    document.add_paragraph("Board meeting minutes for Q3.")
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    buffer.name = "notes.docx"

    text = DocumentProcessor.extract_text(buffer)

    assert "Board meeting minutes for Q3." in text
