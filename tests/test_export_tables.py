"""Tests that markdown tables are included in PDF and DOCX exports."""

from __future__ import annotations

from io import BytesIO

from docx import Document
from PyPDF2 import PdfReader

from services.export_service import ExportService

REPORT_WITH_TABLE = """
## Document Comparison

### Key Differences

| Topic | Document A | Document B |
| --- | --- | --- |
| Revenue | $12.4m | $14.1m |
| Headcount | 210 | 245 |
| Outlook | Stable | Expanding |

Closing note on the comparison.
"""


def test_export_pdf_includes_markdown_tables():
    result = ExportService().export_pdf(
        project_id="export-table-project",
        report_name="Document Comparison",
        report_text=REPORT_WITH_TABLE,
    )

    assert result["data"].startswith(b"%PDF")

    pdf_text = "\n".join(
        page.extract_text() or ""
        for page in PdfReader(BytesIO(result["data"])).pages
    )

    assert "Revenue" in pdf_text
    assert "$12.4m" in pdf_text
    assert "Headcount" in pdf_text
    assert "Expanding" in pdf_text


def test_export_docx_includes_markdown_tables():
    result = ExportService().export_docx(
        project_id="export-table-project",
        report_name="Document Comparison",
        report_text=REPORT_WITH_TABLE,
    )

    assert result["data"].startswith(b"PK")

    document = Document(BytesIO(result["data"]))
    assert document.tables, "expected at least one table in the DOCX export"

    cell_text = [
        cell.text.strip()
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    ]

    assert "Topic" in cell_text
    assert "Document A" in cell_text
    assert "Revenue" in cell_text
    assert "$12.4m" in cell_text
    assert "Expanding" in cell_text
