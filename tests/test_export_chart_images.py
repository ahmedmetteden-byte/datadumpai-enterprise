"""Tests for chart image generation in PDF and Word exports."""

from __future__ import annotations

from io import BytesIO

import pytest
from docx import Document

from services.export_service import ExportService
from services.report_chart_export import (
    CHART_EXPORT_UNAVAILABLE_NOTE,
    is_chart_export_available,
    render_chart_pngs,
)

MINIMAL_PNG = (
    b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01"
    b"\x08\x06\x00\x00\x00\x1f\x15\xc4\x89\x00\x00\x00\nIDATx\x9cc\x00"
    b"\x01\x00\x00\x05\x00\x01\r\n-\xdb\x00\x00\x00\x00IEND\xaeB`\x82"
)


@pytest.fixture
def stub_chart_png_export(monkeypatch):
    is_chart_export_available.cache_clear()
    monkeypatch.setattr(
        "services.report_chart_export.is_chart_export_available",
        lambda: True,
    )
    monkeypatch.setattr(
        "services.report_chart_export.plotly_figure_to_png",
        lambda *args, **kwargs: MINIMAL_PNG,
    )
    yield
    is_chart_export_available.cache_clear()

SAMPLE_CHART_DATA = {
    "topics": [
        {"label": "Claims", "value": 31},
        {"label": "Capital", "value": 21},
    ],
    "trends": [
        {"label": "Claims", "prior": 15, "current": 31},
        {"label": "Capital", "prior": 20, "current": 21},
    ],
    "health_score": 75,
}

FULL_REPORT = """
## Full Report Overview

### Executive Summary Card
| Field | Value |
| Reporting Period | Q1 2024 |
| Confidence | 88% |

## Visual Summary
Charts are rendered by the application.

<!-- REPORT_CHARTS
{
  "topics": [{"label": "Claims", "value": 31}, {"label": "Capital", "value": 21}],
  "trends": [{"label": "Claims", "prior": 15, "current": 31}],
  "health_score": 75
}
-->
"""


def test_render_chart_pngs_returns_png_bytes():
    is_chart_export_available.cache_clear()
    if not is_chart_export_available():
        pytest.skip("Chart export runtime is not available in this environment")

    result = render_chart_pngs(SAMPLE_CHART_DATA)

    assert len(result.images) >= 3
    assert result.unavailable_note is None
    assert result.images[0][0] == "Top Discussion Topics"
    assert result.images[1][0] == "Theme Distribution"
    assert result.images[0][1].startswith(b"\x89PNG")


def test_export_docx_places_charts_near_top_of_document(stub_chart_png_export):
    result = ExportService().export_docx(
        project_id="chart-export-project",
        report_name="Full Report",
        report_text=FULL_REPORT,
    )

    document = Document(BytesIO(result["data"]))
    image_indexes = [
        index
        for index, paragraph in enumerate(document.paragraphs)
        if paragraph._element.xpath(".//a:blip")
    ]
    overview_index = next(
        index
        for index, paragraph in enumerate(document.paragraphs)
        if paragraph.text == "Full Report Overview"
    )

    assert image_indexes
    assert min(image_indexes) < overview_index
    assert any(paragraph.text == "Visual Analytics" for paragraph in document.paragraphs)


def test_export_pdf_embeds_chart_images(stub_chart_png_export):
    result = ExportService().export_pdf(
        project_id="chart-export-project",
        report_name="Full Report",
        report_text=FULL_REPORT,
    )

    assert result["data"].startswith(b"%PDF")
    assert len(result["data"]) > 2000


def test_export_docx_embeds_chart_images(stub_chart_png_export):
    result = ExportService().export_docx(
        project_id="chart-export-project",
        report_name="Full Report",
        report_text=FULL_REPORT,
    )

    assert result["data"].startswith(b"PK")

    document = Document(BytesIO(result["data"]))
    assert any("image" in rel.target_ref for rel in document.part.rels.values())


def test_render_chart_pngs_returns_note_when_export_unavailable(monkeypatch):
    is_chart_export_available.cache_clear()
    monkeypatch.setattr(
        "services.report_chart_export.is_chart_export_available",
        lambda: False,
    )

    result = render_chart_pngs(SAMPLE_CHART_DATA)

    assert result.images == []
    assert result.unavailable_note == CHART_EXPORT_UNAVAILABLE_NOTE

    is_chart_export_available.cache_clear()


def test_export_pdf_continues_without_charts_when_export_unavailable(monkeypatch):
    is_chart_export_available.cache_clear()
    monkeypatch.setattr(
        "services.report_chart_export.is_chart_export_available",
        lambda: False,
    )

    result = ExportService().export_pdf(
        project_id="chart-export-project",
        report_name="Full Report",
        report_text=FULL_REPORT,
    )

    assert result["data"].startswith(b"%PDF")

    is_chart_export_available.cache_clear()


def test_export_docx_includes_chart_unavailable_note(monkeypatch):
    is_chart_export_available.cache_clear()
    monkeypatch.setattr(
        "services.report_chart_export.is_chart_export_available",
        lambda: False,
    )

    result = ExportService().export_docx(
        project_id="chart-export-project",
        report_name="Full Report",
        report_text=FULL_REPORT,
    )

    document = Document(BytesIO(result["data"]))
    assert any(
        paragraph.text == CHART_EXPORT_UNAVAILABLE_NOTE
        for paragraph in document.paragraphs
    )
    assert not any(
        paragraph._element.xpath(".//a:blip")
        for paragraph in document.paragraphs
    )

    is_chart_export_available.cache_clear()
