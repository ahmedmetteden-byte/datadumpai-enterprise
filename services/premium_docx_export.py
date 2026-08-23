"""
Consulting-grade Word export for DataDumpAI executive intelligence reports.
"""

from __future__ import annotations

import re
from dataclasses import dataclass
from datetime import datetime, timezone
from io import BytesIO
from typing import Any

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from config import APP_NAME
from services.export_chart_blocks import get_export_chart_images
from services.report_document_parser import (
    ai_insight_bullets,
    dashboard_metrics,
    estimated_reading_minutes,
    first_ai_insight,
    is_available,
    parse_intelligence_report,
    table_of_contents,
    top_opportunities,
    top_risks,
)
from services.report_markdown_renderer import (
    MarkdownBlock,
    NUMBERED_ITEM_LABEL_PATTERN,
    classify_label_value,
    format_bullet_item,
    group_blocks_for_keep_together,
    parse_markdown_blocks,
    strip_inline_markdown,
)

COLOR_BLUE = RGBColor(0x1D, 0x4E, 0xD8)
COLOR_SLATE = RGBColor(0x0F, 0x17, 0x2A)
COLOR_MUTED = RGBColor(0x64, 0x74, 0x8B)
COLOR_GREEN = RGBColor(0x05, 0x96, 0x69)


@dataclass
class DocxExportMetadata:
    project_name: str
    report_name: str
    report_type: str = "Executive Intelligence Report"
    reporting_period: str = "Not specified"
    source_documents: list[str] | None = None
    pack_type: str = "executive"


def _set_margins(document: Document) -> None:
    for section in document.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)


def _set_paragraph_spacing(paragraph, *, before: int = 6, after: int = 10) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph.paragraph_format.line_spacing = 1.45


def _shade_cell(cell, fill_hex: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def _add_heading(document: Document, text: str, level: int) -> None:
    sizes = {1: 24, 2: 18, 3: 15, 4: 13}
    paragraph = document.add_heading(strip_inline_markdown(text), level=min(level, 4))
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run(strip_inline_markdown(text))
    run.font.name = "Calibri"
    run.font.size = Pt(sizes.get(level, 13))
    run.font.bold = True
    run.font.color.rgb = COLOR_BLUE if level == 1 else COLOR_SLATE
    _set_paragraph_spacing(paragraph, before=6, after=10)
    # A heading must never land alone at the bottom of a page with its
    # content starting fresh on the next — Word's own page-break-avoidance
    # rule for exactly this ("keep with next paragraph").
    paragraph.paragraph_format.keep_with_next = True


def _add_body_paragraph(document: Document, text: str, *, justify: bool = True) -> None:
    paragraph = document.add_paragraph(strip_inline_markdown(text))
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    for run in paragraph.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.font.color.rgb = COLOR_SLATE
    _set_paragraph_spacing(paragraph)


def _rgb_from_hex(hex_color: str) -> RGBColor:
    hex_color = hex_color.lstrip("#")
    return RGBColor(int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16))


def _add_label_value(document: Document, label: str, value: str) -> None:
    """Render a label/value metadata pair (evidence tags, recommendation
    Rationale/Measurement) at a smaller, visually subordinate size than
    body paragraphs — this is metadata about a finding, not the finding
    itself. Uses classify_label_value() as the single source of truth for
    semantic coloring and Source-filename humanization, shared with the
    PDF renderer."""

    label_clean, value_clean, color_hex = classify_label_value(label, value)

    if label_clean.lower() == "basis":
        caption = document.add_paragraph()
        caption_run = caption.add_run("EVIDENCE")
        caption_run.bold = True
        caption_run.font.name = "Calibri"
        caption_run.font.size = Pt(8)
        caption_run.font.color.rgb = COLOR_MUTED
        _set_paragraph_spacing(caption, before=8, after=1)

    label_paragraph = document.add_paragraph()
    label_run = label_paragraph.add_run(f"{label_clean}:")
    label_run.bold = True
    label_run.font.name = "Calibri"
    label_run.font.size = Pt(9)
    label_run.font.color.rgb = COLOR_MUTED
    _set_paragraph_spacing(label_paragraph, before=2, after=1)

    if value_clean:
        value_paragraph = document.add_paragraph(value_clean)
        value_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        value_run = value_paragraph.runs[0]
        value_run.bold = True
        value_run.font.name = "Calibri"
        value_run.font.size = Pt(9)
        value_run.font.color.rgb = _rgb_from_hex(color_hex)
        _set_paragraph_spacing(value_paragraph, before=0, after=10)


def _add_numbered_item(document: Document, text: str) -> None:
    """Render a numbered recommendation item (e.g. its Action clause) at
    normal body size — unlike _add_label_value's evidence tags, this is
    primary content (the recommendation title), not subordinate metadata,
    so it must not silently vanish the way it did before this branch
    existed at all."""

    match = NUMBERED_ITEM_LABEL_PATTERN.match(text)

    if not match:
        _add_body_paragraph(document, text)
        return

    label_clean, value_clean, color_hex = classify_label_value(match.group(1), match.group(2))
    paragraph = document.add_paragraph()
    label_run = paragraph.add_run(f"{label_clean}: ")
    label_run.bold = True
    label_run.font.name = "Calibri"
    label_run.font.size = Pt(11)
    label_run.font.color.rgb = COLOR_SLATE
    value_run = paragraph.add_run(value_clean)
    value_run.bold = True
    value_run.font.name = "Calibri"
    value_run.font.size = Pt(11)
    value_run.font.color.rgb = _rgb_from_hex(color_hex)
    _set_paragraph_spacing(paragraph, before=6, after=4)


def _add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(format_bullet_item(text), style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.35)
    for run in paragraph.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.font.color.rgb = COLOR_SLATE
    _set_paragraph_spacing(paragraph, before=2, after=4)


def _add_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return

    col_count = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"

    for row_index, row in enumerate(rows):
        for col_index in range(col_count):
            cell = table.rows[row_index].cells[col_index]
            cell.text = row[col_index] if col_index < len(row) else ""

            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(11)

            if row_index == 0:
                _shade_cell(cell, "1D4ED8")
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            elif row_index % 2 == 0:
                _shade_cell(cell, "F8FAFC")


def _add_callout(document: Document, title: str, lines: list[str]) -> None:
    table = document.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    _shade_cell(cell, "EFF6FF")

    title_paragraph = cell.paragraphs[0]
    title_run = title_paragraph.add_run(f"💡 {strip_inline_markdown(title)}")
    title_run.bold = True
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(12)
    title_run.font.color.rgb = COLOR_BLUE

    for line in lines:
        body_paragraph = cell.add_paragraph(strip_inline_markdown(line))
        body_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in body_paragraph.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(11)
            run.font.color.rgb = COLOR_SLATE
        _set_paragraph_spacing(body_paragraph)


def _render_blocks(document: Document, blocks: list[MarkdownBlock]) -> None:
    for block in blocks:
        if block.block_type == "heading":
            _add_heading(document, block.content, block.level)
        elif block.block_type == "label_value":
            _add_label_value(document, block.label, block.value)
        elif block.block_type == "bullets":
            for item in block.items:
                _add_bullet(document, item)
        elif block.block_type == "numbered":
            for item in block.items:
                _add_numbered_item(document, item)
        elif block.block_type == "quote":
            paragraph = document.add_paragraph(f"“{block.content}”")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in paragraph.runs:
                run.italic = True
                run.font.name = "Calibri"
                run.font.size = Pt(11)
            _set_paragraph_spacing(paragraph)
        elif block.block_type == "table":
            _add_table(document, block.rows)
        elif block.block_type == "paragraph":
            _add_body_paragraph(document, block.content)


def _append_chart_images(document: Document, chart_data: dict[str, Any]) -> None:
    chart_export = get_export_chart_images(chart_data)

    if not chart_export.images and not chart_export.unavailable_note:
        return

    if chart_export.images:
        for title, png_bytes in chart_export.images:
            _add_heading(document, title, 2)
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.add_run().add_picture(BytesIO(png_bytes), width=Inches(6.2))
        return

    if chart_export.unavailable_note:
        _add_body_paragraph(document, chart_export.unavailable_note)


def _cover_page(document: Document, metadata: DocxExportMetadata) -> None:
    assets_dir = Path(__file__).resolve().parent.parent / "assets"
    logo_path = assets_dir / "logo.png"
    if not logo_path.is_file():
        logo_path = assets_dir / "datadump-hero-logo.png"

    if logo_path.is_file():
        logo_paragraph = document.add_paragraph()
        logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_paragraph.add_run().add_picture(str(logo_path), width=Inches(1.4))
        document.add_paragraph()

    subtitle = document.add_paragraph("Executive Intelligence Report")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(16)
    subtitle_run.font.color.rgb = COLOR_BLUE

    prepared_by_label = document.add_paragraph("Prepared by")
    prepared_by_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    prepared_by_label.runs[0].font.size = Pt(10)
    prepared_by_label.runs[0].font.color.rgb = COLOR_MUTED

    prepared_by_value = document.add_paragraph(APP_NAME)
    prepared_by_value.alignment = WD_ALIGN_PARAGRAPH.CENTER
    prepared_by_value.runs[0].bold = True
    prepared_by_value.runs[0].font.size = Pt(13)

    for label, value in [
        ("Project", metadata.project_name),
        ("Reporting Period", metadata.reporting_period),
        ("Generated", datetime.now(timezone.utc).strftime("%d %B %Y")),
    ]:
        label_paragraph = document.add_paragraph(label)
        label_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        label_paragraph.runs[0].font.size = Pt(10)
        label_paragraph.runs[0].font.color.rgb = COLOR_MUTED

        value_paragraph = document.add_paragraph(value)
        value_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        value_paragraph.runs[0].bold = True
        value_paragraph.runs[0].font.size = Pt(13)

    report_name = document.add_paragraph(metadata.report_name)
    report_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    report_name.runs[0].bold = True
    report_name.runs[0].font.size = Pt(13)

    document.add_paragraph()
    disclaimer = document.add_paragraph(
        "DataDumpAI can make mistakes. Please verify important information before "
        "relying on it."
    )
    disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    disclaimer.runs[0].font.size = Pt(9)
    disclaimer.runs[0].font.color.rgb = COLOR_MUTED

    document.add_page_break()


def _executive_summary_page(document: Document, parsed) -> None:
    metrics = dashboard_metrics(parsed)
    _add_heading(document, "Executive Summary", 1)

    if is_available(metrics.get("health_score")):
        health = document.add_paragraph()
        health.add_run("Overall Health\n").bold = True
        health_value = health.add_run(f"🟢 {metrics['health_score']}/100")
        health_value.bold = True
        health_value.font.color.rgb = COLOR_GREEN
        _set_paragraph_spacing(health)

    if is_available(metrics.get("outlook")):
        outlook = document.add_paragraph()
        outlook.add_run("Overall Outlook\n").bold = True
        outlook.add_run(str(metrics["outlook"]))
        _set_paragraph_spacing(outlook)

    _add_heading(document, "Top Risks", 2)
    risks = top_risks(parsed)[:5]
    if risks:
        for card in risks:
            _add_bullet(document, card["title"])
    else:
        _add_body_paragraph(document, "No risks were identified in the evidence reviewed.")

    _add_heading(document, "Top Opportunities", 2)
    opportunities = top_opportunities(parsed)[:5]
    if opportunities:
        for item in opportunities:
            _add_bullet(document, item)
    else:
        _add_body_paragraph(document, "No opportunities were identified in the evidence reviewed.")

    # The full Strategic Recommendations section (with real Action/
    # Rationale/Measurement structure per item) renders later via
    # _render_blocks(); an early singular one-sentence preview here
    # duplicated it with weaker, regex-extracted text — dropped rather
    # than replaced, matching the PDF renderer.
    _append_chart_images(document, parsed.chart_data)
    document.add_page_break()


def _at_a_glance_page(document: Document, parsed, report_text: str) -> None:
    metrics = dashboard_metrics(parsed)
    _add_heading(document, "At a Glance", 1)

    candidate_rows = [
        ("Overall Status", metrics.get("outlook", "—")),
        ("Health Score", f"{metrics['health_score']}/100" if is_available(metrics.get("health_score")) else ""),
        ("Confidence", metrics.get("confidence", "—")),
        ("Documents", metrics.get("documents", "—")),
        ("Critical Risks", metrics.get("key_risks", "—")),
        ("Top Recommendation", metrics.get("priority", "—")),
        ("Estimated Reading Time", f"{estimated_reading_minutes(report_text)} minutes"),
        ("AI Insight", first_ai_insight(parsed)),
    ]
    rows = [[label, value] for label, value in candidate_rows if is_available(value)]
    _add_table(document, rows)
    document.add_page_break()


def build_premium_docx(
    *,
    report_text: str,
    metadata: DocxExportMetadata,
) -> bytes:
    parsed = parse_intelligence_report(
        report_text,
        source_documents=metadata.source_documents,
        pack_type=metadata.pack_type,
    )

    if not parsed.snapshot.get("Reporting period"):
        period_match = re.search(
            r"Reporting period\s*\|\s*([^|\n]+)",
            report_text,
            re.IGNORECASE,
        )

        if period_match:
            metadata.reporting_period = period_match.group(1).strip()

    document = Document()
    _set_margins(document)

    _cover_page(document, metadata)
    _executive_summary_page(document, parsed)
    _at_a_glance_page(document, parsed, report_text)

    _add_heading(document, "Table of Contents", 1)
    for index, entry in enumerate(
        table_of_contents(parsed.sections, include_appendix=bool(parsed.appendix_sections)),
        start=1,
    ):
        _add_body_paragraph(document, f"{index}. {entry}", justify=False)

    document.add_page_break()

    include_evidence = metadata.pack_type == "board_pack"

    for section in parsed.sections:
        if section.title == "Introduction":
            continue

        _add_heading(document, section.title, 1)
        body = section.body

        if not include_evidence and "Key Findings" in section.title:
            body = re.sub(r"\*\*Evidence:\*\*[\s\S]*?(?=\n\*\*|\n#### |\n### |\Z)", "", body)

        if section.title.lower() == "ai insights":
            bullets = ai_insight_bullets(parsed)
            _add_callout(document, "AI Insight", bullets or [strip_inline_markdown(body)])
            continue

        blocks = parse_markdown_blocks(body)
        groups = group_blocks_for_keep_together(blocks)

        for group in groups:
            _render_blocks(document, group)

    if parsed.appendix_sections:
        document.add_page_break()
        _add_heading(document, "Appendix", 1)

        for section in parsed.appendix_sections:
            _add_heading(document, section.title, 2)
            _render_blocks(document, parse_markdown_blocks(section.body))

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
